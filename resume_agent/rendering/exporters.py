"""In-memory resume export adapters."""

from __future__ import annotations

import shutil
import subprocess
from dataclasses import dataclass
from enum import Enum
from io import BytesIO
from pathlib import Path
from tempfile import TemporaryDirectory
from typing import Iterable, Optional

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

from resume_agent.rendering.models import RenderedResume


class RenderFormat(str, Enum):
    HTML = "html"
    MARKDOWN = "md"
    DOCX = "docx"
    PDF = "pdf"


class RenderEngineUnavailable(RuntimeError):
    """Raised when an optional binary rendering engine is not installed."""


@dataclass(frozen=True)
class ExportedFile:
    filename: str
    media_type: str
    content: bytes


class PdfExporter:
    def __init__(
        self,
        browser_candidates: Optional[Iterable[Path | str]] = None,
        *,
        timeout_seconds: float = 120.0,
    ) -> None:
        self.browser_candidates = (
            list(browser_candidates)
            if browser_candidates is not None
            else self._default_candidates()
        )
        self.timeout_seconds = timeout_seconds

    @staticmethod
    def _default_candidates() -> list[Path | str]:
        candidates: list[Path | str] = [
            Path("/Applications/Google Chrome.app/Contents/MacOS/Google Chrome"),
            Path("/Applications/Microsoft Edge.app/Contents/MacOS/Microsoft Edge"),
        ]
        for command in ("google-chrome", "chromium", "chromium-browser", "msedge"):
            resolved = shutil.which(command)
            if resolved:
                candidates.append(resolved)
        return candidates

    def export(self, html: str) -> bytes:
        available = [Path(item) for item in self.browser_candidates if Path(item).is_file()]
        if not available:
            raise RenderEngineUnavailable(
                "PDF browser engine unavailable; install Google Chrome or Microsoft Edge"
            )
        with TemporaryDirectory(prefix="resume-agent-pdf-") as directory:
            temp_dir = Path(directory)
            html_path = temp_dir / "resume.html"
            pdf_path = temp_dir / "resume.pdf"
            html_path.write_text(html, encoding="utf-8")
            errors = []
            for browser in available:
                try:
                    result = subprocess.run(
                        [
                            str(browser),
                            "--headless=new",
                            "--disable-gpu",
                            "--no-pdf-header-footer",
                            f"--print-to-pdf={pdf_path}",
                            html_path.resolve().as_uri(),
                        ],
                        capture_output=True,
                        check=False,
                        timeout=self.timeout_seconds,
                    )
                except (OSError, subprocess.SubprocessError) as error:
                    errors.append(type(error).__name__)
                    continue
                if result.returncode == 0 and pdf_path.is_file():
                    content = pdf_path.read_bytes()
                    if content.startswith(b"%PDF-"):
                        return content
                errors.append(f"exit {result.returncode}")
            detail = ", ".join(errors) or "no usable browser"
            raise RenderEngineUnavailable(f"PDF browser engine failed: {detail}")


class ResumeExporter:
    def __init__(self, pdf_exporter: Optional[PdfExporter] = None) -> None:
        self.pdf_exporter = pdf_exporter or PdfExporter()

    def export(
        self,
        rendered: RenderedResume,
        format: RenderFormat,
    ) -> ExportedFile:
        if format is RenderFormat.HTML:
            return ExportedFile(
                filename=f"{rendered.filename_stem}.html",
                media_type="text/html",
                content=rendered.html.encode("utf-8"),
            )
        if format is RenderFormat.MARKDOWN:
            return ExportedFile(
                filename=f"{rendered.filename_stem}.md",
                media_type="text/markdown",
                content=rendered.markdown.encode("utf-8"),
            )
        if format is RenderFormat.DOCX:
            return ExportedFile(
                filename=f"{rendered.filename_stem}.docx",
                media_type=(
                    "application/vnd.openxmlformats-officedocument."
                    "wordprocessingml.document"
                ),
                content=self._docx(rendered),
            )
        if format is RenderFormat.PDF:
            return ExportedFile(
                filename=f"{rendered.filename_stem}.pdf",
                media_type="application/pdf",
                content=self.pdf_exporter.export(rendered.html),
            )
        raise ValueError(f"unsupported render format: {format}")

    @staticmethod
    def _docx(rendered: RenderedResume) -> bytes:
        document = Document()
        section = document.sections[0]
        section.top_margin = section.bottom_margin
        title = document.add_heading(rendered.candidate_name, level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.LEFT
        if rendered.headline:
            document.add_paragraph(rendered.headline)
        if rendered.contact_line:
            document.add_paragraph(rendered.contact_line)
        document.add_heading(
            {"zh": "职业概述", "en": "Summary", "ja": "職務要約"}[rendered.locale],
            level=1,
        )
        document.add_paragraph(rendered.summary)
        document.add_heading(
            {"zh": "工作经历", "en": "Experience", "ja": "職務経歴"}[rendered.locale],
            level=1,
        )
        for experience in rendered.experiences:
            heading = f"{experience.role} — {experience.organization}"
            if experience.period:
                heading += f" | {experience.period}"
            document.add_heading(heading, level=2)
            for bullet in experience.bullets:
                document.add_paragraph(bullet, style="List Bullet")
        if rendered.skills:
            document.add_heading(
                {"zh": "技能", "en": "Skills", "ja": "活かせるスキル"}[
                    rendered.locale
                ],
                level=1,
            )
            document.add_paragraph(" · ".join(rendered.skills))
        buffer = BytesIO()
        document.save(buffer)
        return buffer.getvalue()
